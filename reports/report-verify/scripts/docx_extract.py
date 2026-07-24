from __future__ import annotations

import re
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET
from zipfile import ZipFile

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from .utils import normalize


W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
Q = lambda tag: f"{{{W}}}{tag}"


def _cell_text(cell: Any) -> str:
    return "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())


def _table_rows(table: Table) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table.rows:
        values: list[str] = []
        seen_xml_cells: set[int] = set()
        for cell in row.cells:
            identity = id(cell._tc)
            values.append("" if identity in seen_xml_cells else _cell_text(cell))
            seen_xml_cells.add(identity)
        rows.append(values)
    return rows


def _iter_body(document: Document):
    paragraph_by_element = {p._p: p for p in document.paragraphs}
    table_by_element = {t._tbl: t for t in document.tables}
    for child in document.element.body.iterchildren():
        if child in paragraph_by_element:
            yield "paragraph", paragraph_by_element[child]
        elif child in table_by_element:
            yield "table", table_by_element[child]


def extract_comments(path: Path) -> list[dict[str, Any]]:
    with ZipFile(path) as archive:
        if "word/comments.xml" not in archive.namelist():
            return []
        comments_root = ET.fromstring(archive.read("word/comments.xml"))
        document_root = ET.fromstring(archive.read("word/document.xml"))

    comments: dict[str, dict[str, Any]] = {}
    order: list[str] = []
    for element in comments_root.iter(Q("comment")):
        comment_id = element.attrib[Q("id")]
        order.append(comment_id)
        comments[comment_id] = {
            "id": int(comment_id),
            "author": element.attrib.get(Q("author"), ""),
            "date": element.attrib.get(Q("date"), ""),
            "text": "".join(node.text or "" for node in element.iter(Q("t"))),
            "anchor": "",
            "context": "",
        }

    parent = {child: node for node in document_root.iter() for child in node}
    active: list[str] = []
    anchors: dict[str, list[str]] = {key: [] for key in comments}
    contexts: dict[str, str] = {}
    for element in document_root.iter():
        if element.tag == Q("commentRangeStart"):
            comment_id = element.attrib[Q("id")]
            active.append(comment_id)
            node = parent.get(element)
            while node is not None and node.tag not in {Q("p"), Q("tc")}:
                node = parent.get(node)
            if node is not None:
                contexts[comment_id] = "".join(t.text or "" for t in node.iter(Q("t")))
        elif element.tag == Q("commentRangeEnd"):
            comment_id = element.attrib[Q("id")]
            if comment_id in active:
                active.remove(comment_id)
        elif element.tag in {Q("t"), Q("delText")} and element.text:
            for comment_id in active:
                anchors.setdefault(comment_id, []).append(element.text)

    for comment_id, comment in comments.items():
        comment["anchor"] = "".join(anchors.get(comment_id, []))
        comment["context"] = contexts.get(comment_id, comment["anchor"])
    return [comments[comment_id] for comment_id in order]


def extract_docx(path: Path, include_comments: bool = False) -> dict[str, Any]:
    document = Document(path)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    tables: list[dict[str, Any]] = []
    recent_caption = ""
    table_index = 0
    for kind, item in _iter_body(document):
        if kind == "paragraph":
            text = item.text.strip()
            if text and (re.search(r"表\s*\d", normalize(text)) or "结果报告单" in text):
                recent_caption = text
            continue
        rows = _table_rows(item)
        tables.append(
            {
                "index": table_index,
                "caption": recent_caption,
                "row_count": len(rows),
                "rows": rows,
            }
        )
        table_index += 1
        recent_caption = ""

    result: dict[str, Any] = {
        "source": str(path),
        "paragraphs": paragraphs,
        "tables": tables,
    }
    if include_comments:
        result["comments"] = extract_comments(path)
    return result

