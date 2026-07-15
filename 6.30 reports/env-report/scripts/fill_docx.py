#!/usr/bin/env python3
"""Fill the sampling plan template with python-docx while preserving formatting."""

from __future__ import annotations

import argparse
import copy
import json
import sys
from pathlib import Path
from typing import Dict, List, Tuple

from docx.oxml.ns import qn
from docx.shared import Cm, Pt


TABLE2_COLUMNS = ["检测项目", "检测依据", "样品保存条件和期限"]
TABLE3_KEYS = [
    "sampling_no",
    "workplace",
    "position",
    "people_per_shift",
    "job_type",
    "target",
    "project",
    "limit_type",
    "exposure_type",
    "sampling_mode",
    "time_type",
    "collector",
    "device",
    "flow_rate",
    "points_per_day",
    "times_per_day",
    "days",
    "sampling_time",
    "representative_time",
]

DOCUMENT_RULES: Dict[str, object] = {}


def import_docx():
    try:
        from docx import Document
    except ModuleNotFoundError as exc:  # pragma: no cover - runtime env check
        raise RuntimeError("python-docx is not installed") from exc
    return Document


def load_document_rules(path: Path) -> Dict[str, object]:
    if not path.exists():
        raise RuntimeError(f"report rules file not found: {path}")
    rules = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(rules, dict) or not isinstance(rules.get("document"), dict):
        raise RuntimeError("report rules file must contain a document object")
    return rules["document"]


def format_fill_run(run) -> None:
    """Apply the required five-point font to generated form values."""
    run.font.name = DOCUMENT_RULES["fill_font_name"]
    run.font.size = Pt(DOCUMENT_RULES["fill_font_size_pt"])
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), DOCUMENT_RULES["fill_font_name"])
    r_fonts.set(qn("w:hAnsi"), DOCUMENT_RULES["fill_font_name"])
    r_fonts.set(qn("w:eastAsia"), DOCUMENT_RULES["fill_east_asia_font"])


def set_cell_text(cell, text: str) -> None:
    text = (text or "").strip()
    paragraphs = cell.paragraphs
    if not paragraphs:
        paragraph = cell.add_paragraph()
        run = paragraph.add_run(text)
        format_fill_run(run)
        return

    first_para = paragraphs[0]
    if first_para.runs:
        first_para.runs[0].text = text
        format_fill_run(first_para.runs[0])
        for run in first_para.runs[1:]:
            run.text = ""
    else:
        run = first_para.add_run(text)
        format_fill_run(run)

    # Cloned and vertically merged template cells can retain blank paragraphs.
    # Remove them so the value has no trailing line breaks that offset centering.
    for paragraph in paragraphs[1:]:
        paragraph._element.getparent().remove(paragraph._element)


def append_cloned_row(table, template_row, insert_before_row_idx: int):
    new_tr = copy.deepcopy(template_row._tr)
    anchor = table.rows[insert_before_row_idx]._tr
    anchor.addprevious(new_tr)
    return table.rows[insert_before_row_idx]


def ensure_rows(
    table,
    start_row_idx: int,
    data_len: int,
    template_row_idx: int,
    note_row_idx: int,
) -> None:
    needed = start_row_idx + data_len
    while note_row_idx < needed:
        append_cloned_row(table, table.rows[template_row_idx], note_row_idx)
        note_row_idx += 1


def load_payload(path: Path) -> Dict[str, object]:
    return json.loads(path.read_text(encoding="utf-8"))


def set_paragraph_text(paragraph, text: str) -> None:
    text = (text or "").strip()
    if paragraph.runs:
        paragraph.runs[0].text = text
        format_fill_run(paragraph.runs[0])
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        run = paragraph.add_run(text)
        format_fill_run(run)


def format_detection_type(value: str) -> str:
    options = DOCUMENT_RULES["detection_type_options"]
    normalized = (value or "").strip()
    parts = []
    for option in options:
        mark = "☑" if option == normalized else "□"
        parts.append(f"{mark}{option}")
    return "  ".join(parts)


def write_doc_meta(doc, header: Dict[str, str]) -> None:
    if len(doc.paragraphs) > 1:
        set_paragraph_text(doc.paragraphs[1], f"检测任务编号：{header.get('detection_task_no', '')}")


def write_header(table, header: Dict[str, str]) -> None:
    mapping: List[Tuple[Tuple[int, int], str]] = [
        ((0, 1), header.get("unit_name", "")),
        ((0, 5), header.get("contact", "")),
        ((1, 1), header.get("address", "")),
        ((1, 5), format_detection_type(header.get("detection_type", ""))),
    ]
    for (row_idx, cell_idx), value in mapping:
        set_cell_text(table.rows[row_idx].cells[cell_idx], value)


def write_table2(table, rows: List[Dict[str, str]]) -> None:
    start_row = 1
    template_row = 1
    note_row = len(table.rows) - 1
    ensure_rows(table, start_row, len(rows), template_row, note_row)
    for offset, row in enumerate(rows):
        target = table.rows[start_row + offset]
        set_cell_text(target.cells[0], row.get("检测项目", ""))
        set_cell_text(target.cells[1], row.get("检测依据", ""))
        set_cell_text(target.cells[2], row.get("样品保存条件和期限", ""))


def set_table3_column_widths(table) -> None:
    """Reserve four five-point Chinese characters for the object/location column."""
    target_column = 5
    project_column = 6
    points_per_day_column = 14
    sampling_time_column = 17
    grid_columns = table._tbl.tblGrid.gridCol_lst
    if len(grid_columns) <= sampling_time_column:
        return

    target_width = Cm(DOCUMENT_RULES["table3_widths_cm"]["target"])
    points_per_day_width = Cm(DOCUMENT_RULES["table3_widths_cm"]["points_per_day"])
    old_target_width = grid_columns[target_column].w
    project_width = grid_columns[project_column].w + old_target_width - target_width
    old_points_per_day_width = grid_columns[points_per_day_column].w
    sampling_time_width = (
        grid_columns[sampling_time_column].w
        + old_points_per_day_width
        - points_per_day_width
    )
    table.autofit = False
    grid_columns[target_column].w = target_width
    grid_columns[project_column].w = project_width
    grid_columns[points_per_day_column].w = points_per_day_width
    grid_columns[sampling_time_column].w = sampling_time_width
    for row in table.rows:
        row.cells[target_column].width = target_width
        row.cells[project_column].width = project_width
        row.cells[points_per_day_column].width = points_per_day_width
        row.cells[sampling_time_column].width = sampling_time_width


def merge_vertical_same_value_cells(
    table,
    start_row: int,
    end_row: int,
    column_idx: int,
    required_match_column_idxs: Tuple[int, ...] = (),
) -> None:
    if end_row <= start_row:
        return

    group_start = start_row
    current_value = table.rows[start_row].cells[column_idx].text.strip()
    required_match_values = tuple(
        table.rows[start_row].cells[column].text.strip()
        for column in required_match_column_idxs
    )

    for row_idx in range(start_row + 1, end_row + 1):
        row_value = table.rows[row_idx].cells[column_idx].text.strip()
        row_required_match_values = tuple(
            table.rows[row_idx].cells[column].text.strip()
            for column in required_match_column_idxs
        )
        if row_value == current_value and row_required_match_values == required_match_values:
            continue
        if current_value and group_start < row_idx - 1:
            merged_cell = table.cell(group_start, column_idx).merge(table.cell(row_idx - 1, column_idx))
            set_cell_text(merged_cell, current_value)
        group_start = row_idx
        current_value = row_value
        required_match_values = row_required_match_values

    if current_value and group_start < end_row:
        merged_cell = table.cell(group_start, column_idx).merge(table.cell(end_row, column_idx))
        set_cell_text(merged_cell, current_value)


def write_table3(table, rows: List[Dict[str, str]]) -> None:
    set_table3_column_widths(table)
    start_row = 1
    template_row = 1
    note_row = len(table.rows) - 1
    ensure_rows(table, start_row, len(rows), template_row, note_row)
    for offset, row in enumerate(rows):
        target = table.rows[start_row + offset]
        for idx, key in enumerate(TABLE3_KEYS):
            set_cell_text(target.cells[idx], row.get(key, ""))
    if rows:
        end_row = start_row + len(rows) - 1
        # A workplace may span multiple jobs, but must not merge across job boundaries.
        merge_vertical_same_value_cells(table, start_row, end_row, 1, required_match_column_idxs=(2,))
        merge_vertical_same_value_cells(table, start_row, end_row, 2, required_match_column_idxs=(1,))
        # Remaining context values merge independently, but only inside the same
        # workplace-and-job group.
        for column in (3, 4, 5, 8):
            merge_vertical_same_value_cells(table, start_row, end_row, column, required_match_column_idxs=(1, 2))


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", required=True, type=Path)
    parser.add_argument("--payload", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--config", type=Path)
    args = parser.parse_args()

    try:
        global DOCUMENT_RULES
        config_path = args.config or Path(__file__).resolve().parent.parent / "knowledge" / "report_rules.json"
        DOCUMENT_RULES = load_document_rules(config_path)
        Document = import_docx()
        payload = load_payload(args.payload)
        doc = Document(str(args.template))
    except Exception as exc:  # pragma: no cover - CLI error path
        print(str(exc), file=sys.stderr)
        return 1

    tables = doc.tables
    if len(tables) < 3:
        print("template does not contain the expected 3 tables", file=sys.stderr)
        return 1

    write_doc_meta(doc, payload.get("header", {}))
    write_header(tables[0], payload.get("header", {}))
    write_table2(tables[1], payload.get("table2", []))
    write_table3(tables[2], payload.get("table3", []))
    doc.save(str(args.output))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
